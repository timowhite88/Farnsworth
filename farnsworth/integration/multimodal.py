"""
Farnsworth Multimodal Processing - Vision, Audio, and Document Processing

Provides multimodal input handling:
- Image analysis and description
- Audio transcription (Whisper)
- Document parsing (PDF, DOCX, etc.)
- Video frame extraction

Novel Features:
- Streaming audio transcription
- Multi-page document understanding
- Image-text interleaving for context
- Automatic modality detection
"""

import asyncio
import base64
import io
import json
import hashlib
from datetime import datetime
from typing import Optional, Any, Union
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path

from loguru import logger


class Modality(Enum):
    """Supported input modalities."""
    TEXT = "text"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    DOCUMENT = "document"


@dataclass
class MultimodalInput:
    """Represents a multimodal input."""
    modality: Modality
    content: Union[str, bytes]
    mime_type: str
    metadata: dict = field(default_factory=dict)

    # Processing results
    text_representation: Optional[str] = None
    embeddings: Optional[list[float]] = None
    processed_at: Optional[datetime] = None


@dataclass
class ProcessingResult:
    """Result of multimodal processing."""
    success: bool
    text: Optional[str] = None
    structured_data: Optional[dict] = None
    error: Optional[str] = None
    processing_time_ms: float = 0.0
    metadata: dict = field(default_factory=dict)


class MultimodalProcessor:
    """
    Multimodal input processing system.

    Handles:
    - Image analysis (CLIP, BLIP, or vision LLMs)
    - Audio transcription (Whisper)
    - Document extraction (PDFs, Office docs)
    - Video processing (frame extraction)
    """

    def __init__(
        self,
        whisper_model: str = "base",
        enable_gpu: bool = True,
        cache_dir: Optional[str] = None,
    ):
        self.whisper_model_name = whisper_model
        self.enable_gpu = enable_gpu
        self.cache_dir = Path(cache_dir) if cache_dir else None

        # Lazy-loaded models
        self._whisper_model = None
        self._clip_model = None
        self._clip_processor = None

        # Processing cache
        self._cache: dict[str, ProcessingResult] = {}

        # Supported formats
        self.supported_images = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}
        self.supported_audio = {".mp3", ".wav", ".m4a", ".flac", ".ogg", ".wma"}
        self.supported_documents = {".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".rtf"}
        self.supported_video = {".mp4", ".avi", ".mov", ".mkv", ".webm"}

    async def process(
        self,
        input_data: MultimodalInput,
        use_cache: bool = True,
    ) -> ProcessingResult:
        """
        Process a multimodal input.

        Args:
            input_data: Input to process
            use_cache: Whether to use cached results

        Returns:
            ProcessingResult
        """
        # Generate cache key
        if isinstance(input_data.content, bytes):
            cache_key = hashlib.md5(input_data.content).hexdigest()
        else:
            cache_key = hashlib.md5(input_data.content.encode()).hexdigest()

        # Check cache
        if use_cache and cache_key in self._cache:
            return self._cache[cache_key]

        start_time = datetime.now()

        try:
            if input_data.modality == Modality.IMAGE:
                result = await self._process_image(input_data)
            elif input_data.modality == Modality.AUDIO:
                result = await self._process_audio(input_data)
            elif input_data.modality == Modality.DOCUMENT:
                result = await self._process_document(input_data)
            elif input_data.modality == Modality.VIDEO:
                result = await self._process_video(input_data)
            elif input_data.modality == Modality.TEXT:
                result = ProcessingResult(
                    success=True,
                    text=input_data.content if isinstance(input_data.content, str) else input_data.content.decode(),
                )
            else:
                result = ProcessingResult(
                    success=False,
                    error=f"Unsupported modality: {input_data.modality}",
                )

            result.processing_time_ms = (datetime.now() - start_time).total_seconds() * 1000

            # Cache result
            if use_cache and result.success:
                self._cache[cache_key] = result

            return result

        except Exception as e:
            logger.error(f"Multimodal processing failed: {e}")
            return ProcessingResult(
                success=False,
                error=str(e),
                processing_time_ms=(datetime.now() - start_time).total_seconds() * 1000,
            )

    async def _process_image(self, input_data: MultimodalInput) -> ProcessingResult:
        """Process an image input."""
        try:
            # Try to use CLIP or vision model for description
            description = await self._describe_image(input_data.content)

            # Extract any text (OCR)
            ocr_text = await self._ocr_image(input_data.content)

            structured_data = {
                "description": description,
                "ocr_text": ocr_text,
                "dimensions": self._get_image_dimensions(input_data.content),
            }

            # Combine into text representation
            text_parts = [f"[Image: {description}]"]
            if ocr_text:
                text_parts.append(f"[Text in image: {ocr_text}]")

            return ProcessingResult(
                success=True,
                text=" ".join(text_parts),
                structured_data=structured_data,
            )

        except Exception as e:
            return ProcessingResult(success=False, error=str(e))

    async def _process_audio(self, input_data: MultimodalInput) -> ProcessingResult:
        """Process an audio input using Whisper."""
        try:
            # Load Whisper if not loaded
            if self._whisper_model is None:
                self._load_whisper()

            # Transcribe audio
            transcription = await self._transcribe_audio(input_data.content)

            return ProcessingResult(
                success=True,
                text=transcription,
                structured_data={
                    "transcription": transcription,
                    "language": "auto",
                    "duration_seconds": self._get_audio_duration(input_data.content),
                },
            )

        except Exception as e:
            return ProcessingResult(success=False, error=str(e))

    async def _process_document(self, input_data: MultimodalInput) -> ProcessingResult:
        """Process a document input."""
        try:
            mime_type = input_data.mime_type.lower()

            if "pdf" in mime_type:
                text, structured = await self._extract_pdf(input_data.content)
            elif "word" in mime_type or "docx" in mime_type:
                text, structured = await self._extract_docx(input_data.content)
            elif "text" in mime_type or "plain" in mime_type:
                text = input_data.content if isinstance(input_data.content, str) else input_data.content.decode()
                structured = {"type": "plain_text", "length": len(text)}
            else:
                # Try generic text extraction
                text = input_data.content if isinstance(input_data.content, str) else input_data.content.decode("utf-8", errors="ignore")
                structured = {"type": "unknown", "length": len(text)}

            return ProcessingResult(
                success=True,
                text=text,
                structured_data=structured,
            )

        except Exception as e:
            return ProcessingResult(success=False, error=str(e))

    async def _process_video(self, input_data: MultimodalInput) -> ProcessingResult:
        """Process a video input by extracting key frames."""
        try:
            # Extract key frames
            frames = await self._extract_video_frames(input_data.content)

            # Describe each frame
            descriptions = []
            for i, frame in enumerate(frames[:5]):  # Limit to 5 frames
                desc = await self._describe_image(frame)
                descriptions.append(f"Frame {i+1}: {desc}")

            # Transcribe audio if present
            audio_text = await self._extract_video_audio(input_data.content)

            text_parts = ["[Video content:]"]
            text_parts.extend(descriptions)
            if audio_text:
                text_parts.append(f"[Audio: {audio_text}]")

            return ProcessingResult(
                success=True,
                text="\n".join(text_parts),
                structured_data={
                    "frame_descriptions": descriptions,
                    "audio_transcription": audio_text,
                    "frame_count": len(frames),
                },
            )

        except Exception as e:
            return ProcessingResult(success=False, error=str(e))

    def _load_whisper(self):
        """Load Whisper model."""
        try:
            import whisper
            device = "cuda" if self.enable_gpu else "cpu"
            self._whisper_model = whisper.load_model(self.whisper_model_name, device=device)
            logger.info(f"Loaded Whisper model: {self.whisper_model_name}")
        except ImportError:
            logger.warning("Whisper not installed. Audio transcription unavailable.")
            raise RuntimeError("Whisper library not installed. Run: pip install openai-whisper")

    async def _transcribe_audio(self, audio_data: bytes) -> str:
        """Transcribe audio using Whisper."""
        if self._whisper_model is None:
            return "[Audio transcription unavailable - Whisper not loaded]"

        # Save to temp file
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
            f.write(audio_data)
            temp_path = f.name

        try:
            # Run transcription in thread pool
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                lambda: self._whisper_model.transcribe(temp_path),
            )
            return result["text"]
        finally:
            Path(temp_path).unlink(missing_ok=True)

    async def _describe_image(self, image_data: bytes) -> str:
        """Generate a description of an image."""
        # Placeholder - in production, use vision model
        try:
            dimensions = self._get_image_dimensions(image_data)
            return f"An image with dimensions {dimensions.get('width', '?')}x{dimensions.get('height', '?')}"
        except Exception:
            return "An image"

    async def _ocr_image(self, image_data: bytes) -> Optional[str]:
        """Extract text from image using OCR."""
        try:
            # Try pytesseract if available
            import pytesseract
            from PIL import Image

            image = Image.open(io.BytesIO(image_data))
            text = pytesseract.image_to_string(image)
            return text.strip() if text.strip() else None
        except ImportError:
            return None
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return None

    def _get_image_dimensions(self, image_data: bytes) -> dict:
        """Get image dimensions."""
        try:
            from PIL import Image
            image = Image.open(io.BytesIO(image_data))
            return {"width": image.width, "height": image.height, "format": image.format}
        except ImportError:
            # Fallback: try to parse header
            return {}
        except Exception:
            return {}

    def _get_audio_duration(self, audio_data: bytes) -> Optional[float]:
        """Get audio duration in seconds."""
        try:
            import wave
            with wave.open(io.BytesIO(audio_data), "rb") as wav:
                frames = wav.getnframes()
                rate = wav.getframerate()
                return frames / float(rate)
        except Exception:
            return None

    async def _extract_pdf(self, pdf_data: bytes) -> tuple[str, dict]:
        """Extract text and metadata from PDF."""
        try:
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(pdf_data))
            text_parts = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")

            metadata = {
                "type": "pdf",
                "pages": len(reader.pages),
                "info": dict(reader.metadata) if reader.metadata else {},
            }

            return "\n\n".join(text_parts), metadata

        except ImportError:
            # Fallback without pypdf
            return "[PDF extraction requires pypdf library]", {"type": "pdf", "error": "pypdf not installed"}

    async def _extract_docx(self, docx_data: bytes) -> tuple[str, dict]:
        """Extract text from DOCX."""
        try:
            import docx

            doc = docx.Document(io.BytesIO(docx_data))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            metadata = {
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
            }

            return "\n\n".join(text_parts), metadata

        except ImportError:
            return "[DOCX extraction requires python-docx library]", {"type": "docx", "error": "python-docx not installed"}

    async def _extract_video_frames(self, video_data: bytes) -> list[bytes]:
        """Extract key frames from video."""
        try:
            import cv2
            import tempfile

            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix=".mp4", delete=False) as f:
                f.write(video_data)
                temp_path = f.name

            try:
                cap = cv2.VideoCapture(temp_path)
                frames = []
                frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                fps = cap.get(cv2.CAP_PROP_FPS)

                # Extract frames at regular intervals
                interval = max(1, frame_count // 10)

                for i in range(0, frame_count, interval):
                    cap.set(cv2.CAP_PROP_POS_FRAMES, i)
                    ret, frame = cap.read()
                    if ret:
                        _, buffer = cv2.imencode(".jpg", frame)
                        frames.append(buffer.tobytes())

                cap.release()
                return frames

            finally:
                Path(temp_path).unlink(missing_ok=True)

        except ImportError:
            return []

    async def _extract_video_audio(self, video_data: bytes) -> Optional[str]:
        """Extract and transcribe audio from video."""
        try:
            import tempfile
            import subprocess

            # Save video to temp file
            with tempfile.NamedTemporaryFile(suffix=".mp4", delete=False) as vf:
                vf.write(video_data)
                video_path = vf.name

            # Extract audio with ffmpeg
            audio_path = video_path.replace(".mp4", ".wav")

            try:
                subprocess.run([
                    "ffmpeg", "-i", video_path,
                    "-vn", "-acodec", "pcm_s16le",
                    "-ar", "16000", "-ac", "1",
                    audio_path, "-y"
                ], capture_output=True, check=True)

                # Read audio and transcribe
                with open(audio_path, "rb") as af:
                    audio_data = af.read()

                return await self._transcribe_audio(audio_data)

            finally:
                Path(video_path).unlink(missing_ok=True)
                Path(audio_path).unlink(missing_ok=True)

        except Exception as e:
            logger.warning(f"Video audio extraction failed: {e}")
            return None

    def detect_modality(self, file_path: str) -> Modality:
        """
        Detect the modality of a file.

        Args:
            file_path: Path to the file

        Returns:
            Detected Modality
        """
        ext = Path(file_path).suffix.lower()

        if ext in self.supported_images:
            return Modality.IMAGE
        elif ext in self.supported_audio:
            return Modality.AUDIO
        elif ext in self.supported_documents:
            return Modality.DOCUMENT
        elif ext in self.supported_video:
            return Modality.VIDEO
        else:
            return Modality.TEXT

    async def process_file(self, file_path: str) -> ProcessingResult:
        """
        Process a file with automatic modality detection.

        Args:
            file_path: Path to the file

        Returns:
            ProcessingResult
        """
        path = Path(file_path)

        if not path.exists():
            return ProcessingResult(success=False, error=f"File not found: {file_path}")

        modality = self.detect_modality(file_path)

        # Read file content
        with open(path, "rb") as f:
            content = f.read()

        # Determine MIME type
        import mimetypes
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type is None:
            mime_type = "application/octet-stream"

        # Create input
        input_data = MultimodalInput(
            modality=modality,
            content=content,
            mime_type=mime_type,
            metadata={"file_path": file_path, "file_size": len(content)},
        )

        return await self.process(input_data)

    def clear_cache(self) -> int:
        """Clear the processing cache. Returns number of entries cleared."""
        count = len(self._cache)
        self._cache.clear()
        return count


class StreamingAudioProcessor:
    """
    Streaming audio processor for real-time transcription.

    Novel feature: Processes audio in chunks for responsive transcription.
    """

    def __init__(
        self,
        whisper_model: str = "base",
        chunk_duration_seconds: float = 3.0,
        overlap_seconds: float = 0.5,
    ):
        self.whisper_model = whisper_model
        self.chunk_duration = chunk_duration_seconds
        self.overlap = overlap_seconds

        self._model = None
        self._buffer = bytearray()
        self._transcription_buffer = []

    async def initialize(self):
        """Initialize the streaming processor."""
        try:
            import whisper
            self._model = whisper.load_model(self.whisper_model)
            logger.info("Streaming audio processor initialized")
        except ImportError:
            logger.warning("Whisper not available for streaming")

    async def feed_chunk(self, audio_chunk: bytes) -> Optional[str]:
        """
        Feed an audio chunk and get incremental transcription.

        Args:
            audio_chunk: Raw audio bytes

        Returns:
            Transcribed text for this chunk, or None
        """
        self._buffer.extend(audio_chunk)

        # Check if we have enough for processing
        # Assuming 16kHz, 16-bit mono audio
        samples_per_second = 16000 * 2
        min_samples = int(self.chunk_duration * samples_per_second)

        if len(self._buffer) < min_samples:
            return None

        # Process chunk
        chunk = bytes(self._buffer[:min_samples])

        # Keep overlap
        overlap_samples = int(self.overlap * samples_per_second)
        self._buffer = self._buffer[min_samples - overlap_samples:]

        # Transcribe
        if self._model is None:
            return None

        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
            # Write WAV header and data
            self._write_wav(f, chunk)
            temp_path = f.name

        try:
            result = self._model.transcribe(temp_path, fp16=False)
            text = result["text"].strip()

            if text:
                self._transcription_buffer.append(text)
                return text

            return None

        finally:
            Path(temp_path).unlink(missing_ok=True)

    def _write_wav(self, f, data: bytes):
        """Write WAV file with proper headers."""
        import struct

        # WAV header
        sample_rate = 16000
        num_channels = 1
        bits_per_sample = 16
        byte_rate = sample_rate * num_channels * bits_per_sample // 8
        block_align = num_channels * bits_per_sample // 8
        data_size = len(data)

        f.write(b"RIFF")
        f.write(struct.pack("<I", 36 + data_size))
        f.write(b"WAVE")
        f.write(b"fmt ")
        f.write(struct.pack("<IHHIIHH", 16, 1, num_channels, sample_rate, byte_rate, block_align, bits_per_sample))
        f.write(b"data")
        f.write(struct.pack("<I", data_size))
        f.write(data)

    def get_full_transcription(self) -> str:
        """Get the complete transcription so far."""
        return " ".join(self._transcription_buffer)

    def reset(self):
        """Reset the processor state."""
        self._buffer = bytearray()
        self._transcription_buffer = []
